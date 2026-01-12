from io import BytesIO
from typing import Literal, Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors
from PIL import Image
import os
import re
from datetime import datetime

class DocumentGenerator:
    def __init__(self, watermark_path: str = None, watermark_text: str = "FounderGPT"):
        self.watermark_path = watermark_path or os.path.join(
            os.path.dirname(__file__), 
            "../../watermark.png"  # Fallback if frontend path not accessible
        )
        self.watermark_text = watermark_text
        self.company_name = "FounderGPT"
        self.company_url = "https://foundergpt.com"
    
    def generate_pdf(self, content: str, title: str, author: str = None, metadata: Dict[str, Any] = None) -> BytesIO:
        """
        Generate professionally formatted PDF with watermark, proper styling, and metadata.
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Parse markdown content
        parsed_content = self._parse_markdown(content)
        
        # Build story (content elements)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=1,  # Center
            fontName='Helvetica-Bold'
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=20,
            fontName='Helvetica-Bold'
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=10,
            spaceBefore=16,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#333333'),
            spaceAfter=12,
            leading=14,
            fontName='Helvetica'
        )
        
        # Title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Metadata
        if author or metadata:
            meta_text = []
            if author:
                meta_text.append(f"<b>Author:</b> {author}")
            if metadata and metadata.get('document_type'):
                meta_text.append(f"<b>Type:</b> {metadata['document_type']}")
            meta_text.append(f"<b>Generated:</b> {datetime.now().strftime('%B %d, %Y')}")
            
            meta_para = Paragraph(" | ".join(meta_text), styles['Normal'])
            story.append(meta_para)
            story.append(Spacer(1, 0.2*inch))
        
        # Add parsed content
        for element in parsed_content:
            if element['type'] == 'heading1':
                story.append(Paragraph(element['text'], heading1_style))
            elif element['type'] == 'heading2':
                story.append(Paragraph(element['text'], heading2_style))
            elif element['type'] == 'heading3':
                story.append(Paragraph(element['text'], styles['Heading3']))
            elif element['type'] == 'paragraph':
                story.append(Paragraph(element['text'], normal_style))
            elif element['type'] == 'list':
                for item in element['items']:
                    story.append(Paragraph(f"• {item}", normal_style))
                story.append(Spacer(1, 0.1*inch))
            elif element['type'] == 'break':
                story.append(Spacer(1, 0.2*inch))
        
        # Build PDF with watermark
        def add_watermark(canvas_obj, doc_obj):
            self._add_pdf_watermark_advanced(canvas_obj, doc_obj.width, doc_obj.height)
        
        doc.build(story, onFirstPage=add_watermark, onLaterPages=add_watermark)
        
        buffer.seek(0)
        return buffer
    
    def generate_docx(self, content: str, title: str, author: str = None, metadata: Dict[str, Any] = None) -> BytesIO:
        """
        Generate professionally formatted DOCX with watermark, proper styling, and metadata.
        """
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(26, 26, 26)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(12)
        
        # Add metadata
        if author or metadata:
            meta_para = doc.add_paragraph()
            meta_parts = []
            if author:
                meta_parts.append(f"Author: {author}")
            if metadata and metadata.get('document_type'):
                meta_parts.append(f"Type: {metadata['document_type']}")
            meta_parts.append(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            
            meta_run = meta_para.add_run(" | ".join(meta_parts))
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = RGBColor(128, 128, 128)
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_para.space_after = Pt(18)
        
        # Parse and add content
        parsed_content = self._parse_markdown(content)
        
        for element in parsed_content:
            if element['type'] == 'heading1':
                para = doc.add_heading(level=1)
                para.style.font.size = Pt(18)
                para.style.font.color.rgb = RGBColor(44, 62, 80)
                run = para.add_run(element['text'])
                run.bold = True
                para.space_after = Pt(12)
                para.space_before = Pt(20)
            elif element['type'] == 'heading2':
                para = doc.add_heading(level=2)
                para.style.font.size = Pt(14)
                para.style.font.color.rgb = RGBColor(52, 73, 94)
                run = para.add_run(element['text'])
                run.bold = True
                para.space_after = Pt(10)
                para.space_before = Pt(16)
            elif element['type'] == 'heading3':
                para = doc.add_heading(level=3)
                run = para.add_run(element['text'])
                run.bold = True
            elif element['type'] == 'paragraph':
                para = doc.add_paragraph()
                self._add_formatted_text(para, element['text'])
                para.space_after = Pt(12)
            elif element['type'] == 'list':
                for item in element['items']:
                    para = doc.add_paragraph(item, style='List Bullet')
                    para.space_after = Pt(6)
            elif element['type'] == 'break':
                doc.add_paragraph()
        
        # Add watermark
        self._add_docx_watermark_advanced(doc)
        
        # Add footer with company info
        self._add_docx_footer(doc)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _add_pdf_watermark_advanced(self, c: canvas.Canvas, width: float, height: float):
        """
        Add enhanced watermark to PDF page with company branding.
        """
        try:
            c.saveState()
            c.setFillAlpha(0.08)  # Very subtle watermark
            c.setFillColorRGB(0.5, 0.5, 0.5)
            c.setFont("Helvetica-Bold", 72)
            c.translate(width / 2, height / 2)
            c.rotate(45)
            c.drawCentredString(0, 0, self.watermark_text)
            c.restoreState()
            
            # Add footer watermark
            c.saveState()
            c.setFillAlpha(0.3)
            c.setFillColorRGB(0.3, 0.3, 0.3)
            c.setFont("Helvetica", 8)
            footer_text = f"Generated by {self.company_name} - {self.company_url}"
            c.drawString(72, 30, footer_text)
            c.restoreState()
        except Exception as e:
            print(f"Watermark error: {e}")
    
    def _parse_markdown(self, content: str) -> list:
        """
        Parse markdown-like content into structured elements.
        Handles headings, paragraphs, lists, and basic formatting.
        """
        elements = []
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                elements.append({'type': 'break'})
                i += 1
                continue
            
            # Heading 1
            if line.startswith('# '):
                elements.append({
                    'type': 'heading1',
                    'text': line[2:].strip()
                })
            # Heading 2
            elif line.startswith('## '):
                elements.append({
                    'type': 'heading2',
                    'text': line[3:].strip()
                })
            # Heading 3
            elif line.startswith('### '):
                elements.append({
                    'type': 'heading3',
                    'text': line[4:].strip()
                })
            # List item
            elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                items = []
                while i < len(lines) and (lines[i].strip().startswith('- ') or 
                                          lines[i].strip().startswith('* ') or 
                                          lines[i].strip().startswith('• ')):
                    item_text = lines[i].strip()[2:].strip()
                    items.append(item_text)
                    i += 1
                elements.append({
                    'type': 'list',
                    'items': items
                })
                continue
            # Regular paragraph
            else:
                # Collect consecutive non-empty lines as paragraph
                para_lines = []
                while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#'):
                    if not (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                        para_lines.append(lines[i].strip())
                    else:
                        break
                    i += 1
                
                if para_lines:
                    para_text = ' '.join(para_lines)
                    # Basic formatting: **bold**, *italic*
                    para_text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', para_text)
                    para_text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', para_text)
                    elements.append({
                        'type': 'paragraph',
                        'text': para_text
                    })
                continue
            
            i += 1
        
        return elements
    
    def _add_formatted_text(self, para, text: str):
        """
        Add text with basic formatting (bold, italic) to a paragraph.
        """
        # Simple formatting: **bold**, *italic*
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 1:
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part:
                para.add_run(part)
    
    def _add_docx_watermark_advanced(self, doc: Document):
        """
        Add enhanced watermark to DOCX header with company branding.
        """
        try:
            section = doc.sections[0]
            header = section.header
            
            # Clear existing paragraphs
            for para in header.paragraphs:
                para.clear()
            
            # Add watermark text
            header_para = header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = header_para.add_run(self.watermark_text)
            run.font.size = Pt(72)
            run.font.color.rgb = RGBColor(230, 230, 230)  # Very light gray
            run.font.bold = True
            
            # Set paragraph spacing
            header_para.space_after = Pt(0)
        except Exception as e:
            print(f"DOCX watermark error: {e}")
    
    def _add_docx_footer(self, doc: Document):
        """
        Add footer with company information to DOCX.
        """
        try:
            section = doc.sections[0]
            footer = section.footer
            
            footer_para = footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer_text = f"Generated by {self.company_name} - {self.company_url} | {datetime.now().strftime('%B %d, %Y')}"
            run = footer_para.add_run(footer_text)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)
        except Exception as e:
            print(f"DOCX footer error: {e}")

document_generator = DocumentGenerator()
